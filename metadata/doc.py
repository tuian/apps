"""
argv[0] visual map file path with name of file 
argv[1] ifs file path  with name of file 
argv[2] excel sheet path where to create  with name of file 
argv[3] visual map name 
argv[4] xsd service
"""
			
from docx import Document
import os
import xlsxwriter 
import sys  
reload(sys)  
sys.setdefaultencoding('utf8')

workbook = xlsxwriter.Workbook(sys.argv[3])

visualMap = sys.argv[1]
ifsDoc = sys.argv[2]
excel =sys.argv[3]
visualName = sys.argv[4]
xsdName = sys.argv[5]
print sys.argv[1]

worksheet = workbook.add_worksheet("Object Definition")
worksheet2 = workbook.add_worksheet("Mapping Specification")
worksheet2.set_column(0, 8 ,30)
"""worksheet2.set_column('A:A1',15)
worksheet2.set_column('A:B1',15)
worksheet2.set_column('A:C1',30)
worksheet2.set_column('A:D1',30)
worksheet2.set_column('A:E1',15)
worksheet2.set_column('A:F1',15)
worksheet2.set_column('A:G1',30)
worksheet2.set_column('A:H1',30)"""
bold = workbook.add_format({'bold': 1});
worksheet.write('A1', 'System Name', bold);
worksheet.write('B1', 'Instance Name', bold)
worksheet.write('C1', 'Entity Name', bold)
worksheet.write('D1', 'Attribute Name', bold)
worksheet.write('E1', 'Owner', bold)
worksheet.write('F1', 'Parent', bold)
worksheet.write('G1', 'Type', bold)
worksheet.write('H1', 'Description', bold)
worksheet.write('I1', 'URL', bold)
worksheet.write('J1', 'XPATH', bold)
worksheet2.write('A1', 'System Name', bold);
worksheet2.write('B1', 'Instance Name', bold)
worksheet2.write('C1', 'Entity Name', bold)
worksheet2.write('D1', 'Attribute Name', bold)
worksheet2.write('E1', 'System Name', bold);
worksheet2.write('F1', 'Instance Name', bold)
worksheet2.write('G1', 'Entity Name', bold)
worksheet2.write('H1', 'Attribute Name', bold)

file0 = sys.argv[1]
print file0
wordDoc = Document(file0) 
rownum = 1
rowmap=1
col = 0
visualXpath=[]
ifsXpath =[]
ifsAttribute =[]
""" Craete objects for visula map """
for table in wordDoc.tables:
	for row in table.rows:
		if row.cells[1].text.encode("utf-8") != row.cells[3].text.encode("utf-8") :
					worksheet.write_string(rownum ,col , 'Visual Map');
					worksheet.write_string(rownum ,col+1 , 'Default');
					worksheet.write_string(rownum ,col+2 , visualName);
					worksheet.write_string(rownum ,col +3 , row.cells[1].text.encode("utf-8"));
					worksheet.write_string(rownum ,col + 6 , 'UI Field');
					worksheet.write_string(rownum ,col + 7 ,  row.cells[4].text.encode("utf-8"));
					if(row.cells[3].text.encode("utf-8").strip("\n") != 'NA') :
						visualXpath.append(row.cells[3].text.encode("utf-8"))
					"""UX to visual mapping for mapping sheet start """	
					worksheet2.write_string(rowmap ,col , 'UX Sitemap');
					worksheet2.write_string(rowmap ,col+1 , 'Default');
					worksheet2.write_string(rowmap ,col+2 , visualName);
					worksheet2.write_string(rowmap ,col +4 , 'Visual Map');
					worksheet2.write_string(rowmap ,col + 5 , 'Default');
					worksheet2.write_string(rowmap ,col + 6 ,  'Portfolio performance SMA Report');
					worksheet2.write_string(rowmap ,col + 7 ,  row.cells[1].text.encode("utf-8"));
					"""UX to visual mapping for mapping sheet end"""
					rownum = rownum + 1
					rowmap = rowmap + 1					
"""Create onjects for IFS """	
wordDoc = Document(ifsDoc) ; """Ifs argv path"""
addtext = ""
for table in wordDoc.tables:
	for row in table.rows:
		if row.cells[1].text.encode("utf-8") != row.cells[3].text.encode("utf-8") :
				if(row.cells[2].text.encode("utf-8") == "Aggregate") :
					addtext= row.cells[5].text.split("/")[-1];
				worksheet.write_string(rownum ,col , 'IFS');
				worksheet.write_string(rownum ,col + 1 , 'Default');
				worksheet.write_string(rownum ,col + 2 , visualName);
				if(addtext != "") :
					worksheet.write_string(rownum ,col + 3 , addtext +"/"+ row.cells[0].text.encode("utf-8"));
				else :
					worksheet.write_string(rownum ,col + 3 , row.cells[0].text.encode("utf-8"));
				worksheet.write_string(rownum ,col + 6 , 'Attribute');
				worksheet.write_string(rownum ,col + 7 ,  row.cells[1].text.encode("utf-8"));
				rownum = rownum + 1

addtext = ""				
"""Create objects for XSD"""
for table in wordDoc.tables:
	for row in table.rows:
		if row.cells[1].text.encode("utf-8") != row.cells[3].text.encode("utf-8") :
				if(row.cells[2].text.encode("utf-8") == "Aggregate") :
					addtext= row.cells[5].text.split("/")[-1];
				worksheet.write_string(rownum ,col , 'ABS Dev');
				worksheet.write_string(rownum ,col+1 , 'XSD');
				worksheet.write_string(rownum ,col+2 , xsdName);
				if(addtext != "") :
					worksheet.write_string(rownum ,col + 3 , addtext +"/"+ row.cells[0].text.encode("utf-8"));
					ifsAttribute.append(addtext +"/"+ row.cells[0].text.encode("utf-8"));
				else :
					worksheet.write_string(rownum ,col + 3 , row.cells[0].text.encode("utf-8"));
					ifsAttribute.append(row.cells[0].text);
				worksheet.write_string(rownum ,col + 6 , 'Attribute');
				worksheet.write_string(rownum ,col + 7 ,  row.cells[1].text.encode("utf-8"));
				worksheet.write_string(rownum ,col + 9 ,  row.cells[5].text.encode("utf-8"));
				ifsXpath.append(row.cells[5].text.encode("utf-8"))
				rownum = rownum + 1	
"""print ifsXpath
print "\n \n \n "
print ifsAttribute"""

"""Visual map to IFS mapping """
wordDoc= Document(visualMap)
for table in wordDoc.tables:
	for row in table.rows:
		if row.cells[1].text.encode("utf-8") != row.cells[3].text.encode("utf-8") :
					wordDoc1= Document(ifsDoc)
					filedname=""
					if(row.cells[3].text.encode("utf-8").strip("\n") != 'NA') :
						for table1 in wordDoc1.tables:
							for row1 in table1.rows:
								if(row1.cells[2].text.encode("utf-8") == "Aggregate") :
									addtext= row1.cells[5].text.split("/")[-1];
								if(row1.cells[5].text.encode("utf-8").strip() == row.cells[3].text.encode("utf-8").strip()) :
									if(addtext != "") :
										filedname = addtext +"/"+ row1.cells[0].text.encode("utf-8")
									else :
										filedname = row1.cells[0].text.encode("utf-8");
									worksheet2.write_string(rowmap ,col , 'Visual Map');
									worksheet2.write_string(rowmap ,col+1 , 'Default');
									worksheet2.write_string(rowmap ,col+2 , visualName);
									worksheet2.write_string(rowmap ,col+3 ,  row.cells[1].text.encode("utf-8"));
									worksheet2.write_string(rowmap ,col +4 , 'IFS');
									worksheet2.write_string(rowmap ,col + 5 , 'Default');
									worksheet2.write_string(rowmap ,col + 6 , visualName);
									worksheet2.write_string(rowmap ,col + 7 ,filedname);
									rowmap = rowmap + 1					
					
for table in wordDoc.tables:
	for row in table.rows:
		if row.cells[1].text.encode("utf-8") != row.cells[3].text.encode("utf-8") :
					wordDoc1= Document(ifsDoc)
					filedname=""
					if(row.cells[3].text.encode("utf-8").strip("\n") != 'NA') :
						for table1 in wordDoc1.tables:
							for row1 in table1.rows:
								if(row1.cells[2].text.encode("utf-8") == "Aggregate") :
									addtext= row1.cells[5].text.split("/")[-1];
								if(row1.cells[5].text.encode("utf-8").strip() == row.cells[3].text.encode("utf-8").strip()) :
									"""Ifs to xsd"""
						for table1 in wordDoc1.tables:
							for row1 in table1.rows:
								if(row1.cells[2].text.encode("utf-8") == "Aggregate") :
									addtext= row1.cells[5].text.split("/")[-1];
								if(row1.cells[5].text.encode("utf-8").strip() == row.cells[3].text.encode("utf-8").strip()) :
									if(addtext != "") :
										filedname = addtext +"/"+ row1.cells[0].text.encode("utf-8")
									else :
										filedname = row1.cells[0].text.encode("utf-8");
									worksheet2.write_string(rowmap ,col , 'IFS');
									worksheet2.write_string(rowmap ,col+1 , 'Default');
									worksheet2.write_string(rowmap ,col+2 , visualName);
									worksheet2.write_string(rowmap ,col+3 ,  filedname);
									worksheet2.write_string(rowmap ,col +4 , 'ABS Dev');
									worksheet2.write_string(rowmap ,col + 5 , 'XSD');
									worksheet2.write_string(rowmap ,col + 6 , xsdName);
									worksheet2.write_string(rowmap ,col + 7 ,filedname);
									rowmap = rowmap + 1		
	
workbook.close()	