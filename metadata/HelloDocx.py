"""
##########README############
pip install python-docx
"""
from mdr_util import *
import openpyxl
import warnings
import xlsxwriter
import os
from docx import *
import docx
from docx.table import _Cell, Table
from docx.document import Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl

def setHeaderRow_VM_Extraction():

    bold = workbook.add_format({'bold': 1});



    excel_worksheet_screens_objects = workbook.add_worksheet(worksheet_name_screen_objects)
    excel_worksheet_vm_objects = workbook.add_worksheet(worksheet_name_vm_objects)
    excel_worksheet_ifs_objects = workbook.add_worksheet(worksheet_name_ifs_objects)
    excel_worksheet_xsd_objects = workbook.add_worksheet(worksheet_name_xsd_objects)

    excel_worksheet_screen_vm_mappings = workbook.add_worksheet(worksheet_name_screen_vm_mappings)
    excel_worksheet_vm_ifs_mappings = workbook.add_worksheet(worksheet_name_vm_ifs_mappings)
    excel_worksheet_ifs_xsd_mappings = workbook.add_worksheet(worksheet_name_ifs_xsd_mappings)

    excel_worksheet_screens_objects.write('A1', 'System Name', bold)
    excel_worksheet_screens_objects.write('B1', 'Instance Name', bold)
    excel_worksheet_screens_objects.write('C1', 'Entity Name', bold)
    excel_worksheet_screens_objects.write('D1', 'Attribute Name', bold)
    excel_worksheet_screens_objects.write('E1', 'Owner', bold)
    excel_worksheet_screens_objects.write('F1', 'Parent', bold)
    excel_worksheet_screens_objects.write('G1', 'Entity Type', bold)
    excel_worksheet_screens_objects.write('H1', 'Description', bold)
    excel_worksheet_screens_objects.write('I1', 'URL', bold)
    excel_worksheet_screens_objects.write('J1', 'XPATH', bold)
    excel_worksheet_screens_objects.write('K1', 'MDR Phase', bold)

    excel_worksheet_vm_objects.write('A1', 'Visual Map Name', bold)
    excel_worksheet_vm_objects.write('B1', 'Visual Map Attribute Name', bold)
    excel_worksheet_vm_objects.write('C1', 'Description', bold)

    excel_worksheet_ifs_objects.write('A1', 'IFS Name', bold)
    excel_worksheet_ifs_objects.write('B1', 'IFS Attribute Name', bold)
    excel_worksheet_ifs_objects.write('C1', 'Description', bold)

    excel_worksheet_xsd_objects.write('A1', 'XSD Name', bold)
    excel_worksheet_xsd_objects.write('B1', 'XSD Attribute Name', bold)
    excel_worksheet_xsd_objects.write('C1', 'Description', bold)

    excel_worksheet_screen_vm_mappings.write('A1', 'Source System Name', bold)
    excel_worksheet_screen_vm_mappings.write('B1', 'Source Instance Name', bold)
    excel_worksheet_screen_vm_mappings.write('C1', 'Source Entity Name', bold)
    excel_worksheet_screen_vm_mappings.write('D1', 'Source Attribute Name', bold)
    excel_worksheet_screen_vm_mappings.write('E1', 'Target System Name', bold)
    excel_worksheet_screen_vm_mappings.write('F1', 'Target Instance Name', bold)
    excel_worksheet_screen_vm_mappings.write('G1', 'Target Entity Name', bold)
    excel_worksheet_screen_vm_mappings.write('H1', 'Target Attribute Name', bold)
    excel_worksheet_screen_vm_mappings.write('I1', 'Description', bold)

    excel_worksheet_vm_ifs_mappings.write('A1', 'Source System Name', bold)
    excel_worksheet_vm_ifs_mappings.write('B1', 'Source Instance Name', bold)
    excel_worksheet_vm_ifs_mappings.write('C1', 'Source Entity Name', bold)
    excel_worksheet_vm_ifs_mappings.write('D1', 'Source Attribute Name', bold)
    excel_worksheet_vm_ifs_mappings.write('E1', 'Target System Name', bold)
    excel_worksheet_vm_ifs_mappings.write('F1', 'Target Instance Name', bold)
    excel_worksheet_vm_ifs_mappings.write('G1', 'Target Entity Name', bold)
    excel_worksheet_vm_ifs_mappings.write('H1', 'Target Attribute Name', bold)
    excel_worksheet_vm_ifs_mappings.write('I1', 'Description', bold)

    excel_worksheet_ifs_xsd_mappings.write('A1', 'Source System Name', bold)
    excel_worksheet_ifs_xsd_mappings.write('B1', 'Source Instance Name', bold)
    excel_worksheet_ifs_xsd_mappings.write('C1', 'Source Entity Name', bold)
    excel_worksheet_ifs_xsd_mappings.write('D1', 'Source Attribute Name', bold)
    excel_worksheet_ifs_xsd_mappings.write('E1', 'Target System Name', bold)
    excel_worksheet_ifs_xsd_mappings.write('F1', 'Target Instance Name', bold)
    excel_worksheet_ifs_xsd_mappings.write('G1', 'Target Entity Name', bold)
    excel_worksheet_ifs_xsd_mappings.write('H1', 'Target Attribute Name', bold)
    excel_worksheet_ifs_xsd_mappings.write('I1', 'Description', bold)

    excel_worksheet_table_extract.set_column(0, 10, 25)
    excel_worksheet_screens_objects.set_column(0, 10, 25)
    excel_worksheet_vm_objects.set_column(0, 10, 25)
    excel_worksheet_ifs_objects.set_column(0, 10, 25)
    excel_worksheet_xsd_objects.set_column(0, 10, 25)
    excel_worksheet_screen_vm_mappings.set_column(0, 10, 25)
    excel_worksheet_vm_ifs_mappings.set_column(0, 10, 25)
    excel_worksheet_ifs_xsd_mappings.set_column(0, 10, 25)


def iter_block_items(parent):
    #print "Inside"
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        #print "Document"
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        #print "Cell"
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    #parent_elm = parent.element

    for child in parent_elm.iterchildren():
        #print "iterate"
        if isinstance(child, CT_P):
            #print "Paragraph"
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            #print "Table"
            yield Table(child, parent)

file_folder_path = 'C:\VM\\'
file_extension = '.docx'

input_excel_filename = 'VM Input.xlsx'
input_excel_filepath =  file_folder_path + input_excel_filename


warnings.simplefilter("ignore")

#read excel sheet for input
try:

	wb = openpyxl.load_workbook(input_excel_filepath)#open input excel sheet
except IOError:
	print "ERROR: Could not open the input excel spreadsheet."
	print "ERROR DETAILS: Input file '" + input_excel_filepath +"' is missing. Please check the file name and directory details."
else:
	#open sheet by name
	sheet = wb.get_sheet_by_name('Sheet1')


output_excel_folder_path = 'C:\VM\Output\\'

if not os.path.exists(output_excel_folder_path):
    os.makedirs(output_excel_folder_path)

for row in range(2, (sheet.max_row+1)):
#    print "Hello World"
    if ( sheet['A' + str(row)].value != None):
        if((sheet['D' + str(row)].value).lower() == "Yes".lower()):
            print "File Name = ", sheet['A' + str(row)].value

            #filename = 'Visual Map - Fund administration - Subscription'
            filename = sheet['A' + str(row)].value
            filename_complete = file_folder_path + filename + file_extension

            doc = docx.Document(filename_complete)
            #print "Test"


            output_excel_file_name = filename+'.xlsx'

            worksheet_name_table_extract = 'VM Document Table Extract'
            worksheet_name_screen_objects = 'Screen Objects'
            worksheet_name_vm_objects = 'VM Objects'
            worksheet_name_ifs_objects = 'IFS  Objects'
            worksheet_name_xsd_objects = 'XSD  Objects'

            worksheet_name_screen_vm_mappings = 'Screen_VM_Mappings'
            worksheet_name_vm_ifs_mappings = 'VM_IFS_Mappings'
            worksheet_name_ifs_xsd_mappings = 'IFS_XSD_Mappings'

            workbook = xlsxwriter.Workbook(output_excel_folder_path+output_excel_file_name)

            excel_worksheet_table_extract = workbook.add_worksheet(worksheet_name_table_extract)

            setHeaderRow_VM_Extraction()

            row_position = 0

            for block in iter_block_items(doc):
                #print "Class ",block.__class__.__name__
                #print "Test"
                #print("found a block")
                if (block.__class__.__name__ == 'Paragraph'): print block.text

                if (block.__class__.__name__ == 'Table'):

                    table_column_length = len(block.rows[0].cells)
                    table_row_length = len(block.rows)

                    print "table_column_length=",table_column_length
                    print "row count=",table_row_length

                    for row in block.rows:
                        #print len(row.cells)
                        for i in range(table_column_length):
                            excel_worksheet_table_extract.write_string(row_position, i, row.cells[i].text)
                            #excel_worksheet.write_string(row_position, i, "TEST")
                            print row.cells[i].text, " | ",
                        row_position = row_position + 1
                        print "\n"

            workbook.close()

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        return '\n'.join(fullText)
    #return '\n'.join(fullText)
"""
for para in doc.paragraphs:
    if(isinstance(para,Paragraph)):
        print "Paragraph ",para.text

    if(isinstance(para,Table)):
        print "Table ",para.text

"""
#print getText(filename)
#print len(doc.paragraphs)


